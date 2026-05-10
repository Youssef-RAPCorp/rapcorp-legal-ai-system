"""
╔═══════════════════════════════════════════════════════════════════════════════╗
║                    DOCX WRITER & EDITOR UTILITIES                             ║
║              RAPCorp Legal AI System                                          ║
║                                                                               ║
║  Provides:                                                                    ║
║  • txt_to_docx     — converts plain-text legal doc → formatted .docx         ║
║  • find_replace_docx — find/replace text in a .docx file                     ║
║  • count_occurrences — count matches without modifying the file               ║
║  • ai_fix_docx     — apply an AI-generated targeted edit to a .docx          ║
╚═══════════════════════════════════════════════════════════════════════════════╝
"""

import asyncio
import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════════════
# TXT → DOCX CONVERSION
# ═══════════════════════════════════════════════════════════════════════════════

# Tab stop for caption bracket column: 3.5 inches from the left text margin.
# Twips = 1/1440 inch.  3.5 × 1440 = 5040.
_CAPTION_TAB_TWIPS = 5040


def _add_caption_tab_stop(para) -> None:
    """Set a left-aligned tab stop at _CAPTION_TAB_TWIPS on *para*.

    This makes every \t before ) in a caption line jump to exactly the same
    horizontal position regardless of how wide (or bold) the preceding text is.
    """
    pPr = para._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(_CAPTION_TAB_TWIPS))
    tabs_el.append(tab)
    pPr.append(tabs_el)


def txt_to_docx(txt_path: str, docx_path: str, title: str = "") -> None:
    """
    Convert a plain-text legal document to a properly formatted .docx file.

    Formatting heuristics:
      • ALL CAPS lines              → bold centered heading
      • Lines matching /^\\d+\\.\\s/ → numbered paragraph (indented)
      • Lines starting with •/–/-   → bullet item (indented)
      • ═══/───/===/ separator lines → thin horizontal rule paragraph
      • Everything else             → normal body paragraph

    Args:
        txt_path:  Path to the source .txt file.
        docx_path: Path to write the output .docx file.
        title:     Optional document title (added as first heading if provided).
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    try:
        content = Path(txt_path).read_text(encoding="utf-8")
    except OSError as exc:
        raise OSError(f"Could not read source file '{txt_path}': {exc}") from exc
    doc = Document()

    # ── Page layout ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # ── Default style ─────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    if title:
        p = doc.add_paragraph()
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

    _SEPARATOR_RE = re.compile(r'^[═─━=\-─]{4,}$')
    _NUMBERED_RE  = re.compile(r'^\d+[.)]\s')

    in_caption = True   # True until the first === separator line

    for line in content.split("\n"):
        stripped = line.strip()

        # Track end of caption region
        if in_caption and _SEPARATOR_RE.match(stripped):
            in_caption = False

        # Caption bracket line — left text + TAB + ) + optional case info
        # _fix_caption_parens emits "LEFT\t)" or "LEFT\t)    CASE INFO"
        if in_caption and '\t)' in line:
            parts = line.split('\t)', 1)
            left_text = parts[0]
            right_text = parts[1].strip() if len(parts) > 1 else ""
            p = doc.add_paragraph()
            _add_caption_tab_stop(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            # Left party text (may be bold if ALL CAPS — honour that)
            rl = p.add_run(left_text)
            rl.font.size = Pt(12)
            if left_text == left_text.upper() and left_text.strip():
                rl.bold = True
            # Tab then bracket
            p.add_run('\t')
            rb = p.add_run(')')
            rb.font.size = Pt(12)
            # Right-column case info (Case No., Judge, etc.)
            if right_text:
                rc = p.add_run('    ' + right_text)
                rc.font.size = Pt(12)
            continue

        # Separator → thin horizontal rule (em-dashes)
        if _SEPARATOR_RE.match(stripped):
            p = doc.add_paragraph()
            r = p.add_run("─" * 60)
            r.font.size = Pt(8)
            r.font.color.rgb = None   # let Word theme handle it
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            continue

        # Blank line → paragraph break
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            continue

        # ALL CAPS heading (court name, section title, etc.)
        if (stripped == stripped.upper()
                and len(stripped) > 4
                and not stripped.startswith("[")
                and not _NUMBERED_RE.match(stripped)):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.bold = True
            r.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            continue

        # Numbered paragraph
        if _NUMBERED_RE.match(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent   = Inches(0.5)
            p.paragraph_format.space_after   = Pt(4)
            r = p.add_run(stripped)
            r.font.size = Pt(12)
            continue

        # Bullet / checklist item
        if stripped.startswith(("•", "–", "[ ]", "[x]", "[X]")) or (
            stripped.startswith("- ") and len(stripped) > 2
        ):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped)
            r.font.size = Pt(12)
            continue

        # Signature / underline placeholders
        if stripped.startswith("_" * 5) or stripped.startswith("Date:"):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            continue

        # Normal body paragraph
        p = doc.add_paragraph()
        r = p.add_run(stripped)
        r.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    try:
        doc.save(docx_path)
    except OSError as exc:
        raise OSError(f"Could not save '{docx_path}': {exc}") from exc


# ═══════════════════════════════════════════════════════════════════════════════
# FIND / REPLACE
# ═══════════════════════════════════════════════════════════════════════════════

def _iter_paragraphs(doc: "Document"):
    """Yield every paragraph in the doc body and all table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def find_replace_docx(file_path: str, find_text: str, replace_text: str) -> int:
    """
    Find and replace all occurrences of find_text in a .docx file.

    Handles run-split text (Word often splits a word across multiple runs due
    to spell-check or formatting marks) by operating on the full paragraph text
    and rewriting the first run.

    Returns the number of substitutions made.
    """
    if not find_text:
        return 0

    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    if not file_path.endswith(".docx"):
        # Plain-text file — safe to do a text replacement
        try:
            content = Path(file_path).read_text(encoding="utf-8")
            count = content.count(find_text)
            Path(file_path).write_text(
                content.replace(find_text, replace_text), encoding="utf-8"
            )
            return count
        except OSError as exc:
            raise OSError(f"Could not edit '{file_path}': {exc}") from exc

    try:
        doc = Document(file_path)
    except Exception as exc:
        raise ValueError(f"Could not open '{file_path}' as a Word document: {exc}") from exc

    total = 0
    for para in _iter_paragraphs(doc):
        if find_text not in para.text:
            continue
        total += para.text.count(find_text)
        new_text = para.text.replace(find_text, replace_text)
        for i, run in enumerate(para.runs):
            run.text = new_text if i == 0 else ""

    try:
        doc.save(file_path)
    except OSError as exc:
        raise OSError(f"Could not save '{file_path}': {exc}") from exc
    return total


def count_occurrences(file_path: str, find_text: str) -> int:
    """
    Count occurrences of find_text in a .docx or plain-text file.
    Does not modify the file.
    """
    if not find_text:
        return 0

    if DOCX_AVAILABLE and file_path.endswith(".docx"):
        doc = Document(file_path)
        return sum(para.text.count(find_text) for para in _iter_paragraphs(doc))
    else:
        content = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        return content.count(find_text)


# ═══════════════════════════════════════════════════════════════════════════════
# AI-ASSISTED FIX
# ═══════════════════════════════════════════════════════════════════════════════

async def ai_fix_docx(config, file_path: str, instruction: str) -> None:
    """
    Apply a seamlessly integrated AI edit to a legal document.

    Two-stage pipeline:
      Stage 1 (Flash, fast) — analyzes document structure: identifies the
        document type, the exact section/paragraph the change belongs in,
        the paragraph numbering scheme, and any renumbering needed.
      Stage 2 (Pro, full rewrite) — executes the edit using the Stage 1
        plan so the change fits perfectly in context, tone, and numbering.

    Works on both .docx and .txt.  When a .txt source exists for a .docx,
    edits the .txt and reconverts to preserve .docx formatting fidelity.
    """
    try:
        import google.generativeai as genai
        from google.generativeai.types import GenerationConfig
    except ImportError:
        print("  google-generativeai not available — AI fix skipped.")
        return

    p = Path(file_path)

    # Prefer the .txt source when editing a .docx — editing txt then reconverting
    # is more reliable than round-tripping through docx paragraph extraction.
    txt_source = p.with_suffix(".txt")
    if p.suffix.lower() == ".docx" and txt_source.exists():
        content = txt_source.read_text(encoding="utf-8", errors="replace")
        write_txt = str(txt_source)
        write_docx = file_path
    elif p.suffix.lower() == ".docx" and DOCX_AVAILABLE:
        doc = Document(file_path)
        content = "\n".join(p2.text for p2 in _iter_paragraphs(doc))
        write_txt = None
        write_docx = file_path
    else:
        content = p.read_text(encoding="utf-8", errors="replace")
        write_txt = file_path
        write_docx = None

    genai.configure(api_key=config.google_api_key)

    # ── Stage 1: structural analysis (Flash — fast and cheap) ────────────────
    analysis_prompt = f"""You are analyzing a legal document before editing it.

REQUESTED CHANGE:
{instruction}

DOCUMENT:
{content[:8000]}{"[...truncated for analysis...]" if len(content) > 8000 else ""}

Answer the following — be specific and precise:

1. DOCUMENT TYPE: What kind of legal document is this?
   (petition, affidavit, motion, proposed order, exhibit index, etc.)

2. RELEVANT SECTION: Which named section or heading does this change belong under?
   Quote the exact heading text if one exists.

3. INSERTION POINT: After which specific paragraph or sentence should new content
   be inserted? Quote the last sentence of that paragraph verbatim.
   If the change modifies existing text, quote the exact sentence to modify.

4. NUMBERING SCHEME: How are paragraphs numbered?
   (e.g., "¶ 1.", "1.", "(a)", no numbers, Roman numerals — give an example)

5. RENUMBERING NEEDED: If content is inserted mid-document, which paragraph
   numbers must be updated? Give the old → new mapping (e.g., ¶ 14 → ¶ 15).

6. TONE AND STYLE: In 1-2 sentences, describe the writing style of the surrounding
   text so the new content can match it exactly.

7. EDIT PLAN: In 1-2 sentences, state exactly what you will do."""

    flash_model = genai.GenerativeModel(config.model_flash)
    flash_cfg   = GenerationConfig(temperature=0.1, max_output_tokens=1024)
    try:
        analysis_resp = await asyncio.to_thread(
            flash_model.generate_content, analysis_prompt,
            generation_config=flash_cfg,
        )
        analysis = (analysis_resp.text or "").strip()
    except Exception as exc:
        print(f"  Warning: structure analysis failed ({exc}) — proceeding without plan.")
        analysis = "No structural analysis available."

    # ── Stage 2: execute the edit (Pro — full document rewrite) ──────────────
    exec_prompt = f"""You are a senior attorney editing a court filing.
Apply the requested change so that the result reads as though the document
was always written this way — perfectly integrated in placement, tone, and format.

═══ REQUESTED CHANGE ═══
{instruction}

═══ STRUCTURAL ANALYSIS (follow this plan exactly) ═══
{analysis}

═══ EDITING RULES ═══
• Place new content at the EXACT insertion point identified above — not at the end
• Match the paragraph numbering scheme character-for-character (e.g. "¶ 14." not "14.")
• Renumber ALL subsequent paragraphs if inserting mid-document
• Write new sentences in the same legal register as the surrounding text
• Add a transitional phrase where the new paragraph follows from prior context
  (e.g., "Following that incident...", "In addition to the foregoing...",
   "Consistent with the above,...")
• If modifying existing text, change only the targeted sentence(s)
• Preserve every section heading, caption, signature block, and exhibit label exactly
• Do NOT add any new section headers or labels that don't already exist
• Do NOT alter any content not directly affected by the requested change

Return ONLY the final revised document — no analysis, no commentary, no markdown fences.
The document must be clean and ready to file.

═══ ORIGINAL DOCUMENT ═══
{content}

═══ REVISED DOCUMENT ═══"""

    pro_model = genai.GenerativeModel(config.model_pro)
    pro_cfg   = GenerationConfig(temperature=0.15, max_output_tokens=16384)
    try:
        exec_resp = await asyncio.to_thread(
            pro_model.generate_content, exec_prompt, generation_config=pro_cfg,
        )
        revised = (exec_resp.text or "").strip()
    except Exception as exc:
        print(f"  AI fix failed for {p.name}: {exc}")
        return

    if not revised:
        print(f"  AI returned empty output — {p.name} unchanged.")
        return

    # Strip any markdown fences the model may have added
    if revised.startswith("```"):
        revised = re.sub(r'^```[^\n]*\n?', '', revised)
        revised = re.sub(r'\n?```\s*$', '', revised)

    # ── Write results ─────────────────────────────────────────────────────────
    if write_txt:
        Path(write_txt).write_text(revised, encoding="utf-8")
    if write_docx:
        tmp = str(Path(write_docx).with_suffix("._aitmp.txt"))
        Path(tmp).write_text(revised, encoding="utf-8")
        try:
            txt_to_docx(tmp, write_docx)
        finally:
            Path(tmp).unlink(missing_ok=True)
    print(f"  AI fix applied → {p.name}")
