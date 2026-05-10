"""
╔═══════════════════════════════════════════════════════════════════════════════╗
║                    CASE DIRECTORY SCANNER                                     ║
║              RAPCorp Legal AI System                                          ║
║                                                                               ║
║  Reads all legal documents from a directory and extracts their text          ║
║  content for AI-based case analysis.                                          ║
║                                                                               ║
║  Supported formats:                                                           ║
║  • .txt / .md / .rtf / .csv / .json / .html  — read directly                ║
║  • .pdf                                       — extracted via pypdf           ║
║  • .docx / .doc                               — extracted via python-docx    ║
╚═══════════════════════════════════════════════════════════════════════════════╝
"""

import io
import zipfile
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Tuple


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".rtf", ".csv", ".json", ".html", ".htm", ".log",
    ".pdf",
    ".docx", ".doc",
    ".png", ".jpg", ".jpeg", ".gif", ".tiff", ".tif", ".bmp", ".webp",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".tiff", ".tif", ".bmp", ".webp"}

MAX_CHARS_PER_DOC = None  # No truncation — feed the full document
MAX_TOTAL_CHARS   = None  # No total cap — include all documents

# Filename keywords → document type label
_TYPE_KEYWORDS: Dict[str, List[str]] = {
    "petition":       ["petition", "complaint", "original_petition"],
    "response":       ["response", "answer", "reply", "opposition", "objection"],
    "court_order":    ["order", "ruling", "judgment", "decree", "decision", "mandate"],
    "motion":         ["motion"],
    "affidavit":      ["affidavit", "declaration", "sworn", "verification"],
    "summons":        ["summons", "notice_of"],
    "exhibit":        ["exhibit"],
    "checklist":      ["checklist"],
    "cover_sheet":    ["cover_sheet", "cover sheet"],
    "proposed_order": ["proposed_order"],
    "transcript":     ["transcript"],
    "brief":          ["brief", "memorandum", "memo"],
    "evidence_plan":  ["document_plan"],
}


# ═══════════════════════════════════════════════════════════════════════════════
# DATA STRUCTURES
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class CaseDocument:
    """A single document read from the case directory."""
    filename: str
    file_path: str
    doc_type: str           # Inferred: petition, response, court_order, motion, etc.
    text_content: str       # Extracted text (may be truncated)
    char_count: int
    truncated: bool = False
    read_error: Optional[str] = None


# ═══════════════════════════════════════════════════════════════════════════════
# INTERNAL READERS
# ═══════════════════════════════════════════════════════════════════════════════

def _infer_doc_type(filename: str) -> str:
    name = filename.lower().replace(" ", "_").replace("-", "_")
    for doc_type, keywords in _TYPE_KEYWORDS.items():
        if any(kw in name for kw in keywords):
            return doc_type
    return "document"


def _read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()
    # OSError / PermissionError propagate to _read_file, which sets read_error


_SCANNED_PDF_THRESHOLD = 200  # chars; below this, pypdf result is considered insufficient


def _read_pdf(path: str, gemini_api_key: str = "") -> str:
    """
    Extract text from a PDF.

    Strategy:
    1. Try pypdf (fast, free, works on text-layer PDFs).
    2. If pypdf yields fewer than _SCANNED_PDF_THRESHOLD chars (scanned/image PDF),
       fall back to Gemini Flash via the File API, which can OCR scanned documents.
    3. If no API key is available, return whatever pypdf managed to extract.
    """
    pypdf_text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        pypdf_text = "\n\n".join(pages)
    except ImportError:
        pypdf_text = ""
    except Exception:
        pypdf_text = ""

    if len(pypdf_text) >= _SCANNED_PDF_THRESHOLD:
        return pypdf_text  # Good enough — no need for Gemini

    # pypdf got little/nothing — try Gemini File API for scanned PDFs
    if not gemini_api_key:
        return pypdf_text or "[PDF appears to be a scanned image; set GOOGLE_API_KEY to enable OCR extraction]"

    try:
        import google.generativeai as genai
        from google.generativeai.types import GenerationConfig

        genai.configure(api_key=gemini_api_key)

        uploaded = genai.upload_file(path=path, mime_type="application/pdf")

        model = genai.GenerativeModel("gemini-flash-latest")
        response = model.generate_content(
            [
                uploaded,
                "This is a legal document. Extract ALL text exactly as written.\n"
                "- Preserve paragraph structure and section headings.\n"
                "- Extract every table row and column verbatim.\n"
                "- Include all dates, names, case numbers, and signature blocks.\n"
                "- Include headers, footers, and margin annotations.\n"
                "- Do not summarize, paraphrase, or omit anything.\n"
                "Output the complete verbatim text content only.",
            ],
            generation_config=GenerationConfig(temperature=0.0, max_output_tokens=16384),
        )

        extracted = response.text.strip() if hasattr(response, "text") else ""

        # Clean up the uploaded file from Gemini's servers
        try:
            genai.delete_file(uploaded.name)
        except Exception:
            pass

        return extracted if extracted else (pypdf_text or "[Gemini extraction returned empty result]")

    except Exception as e:
        return pypdf_text or f"[PDF extraction failed: {e}]"


def _read_docx(path: str) -> str:
    # Try python-docx first — extracts paragraphs AND tables
    try:
        import docx
        from docx.oxml.ns import qn as _qn

        doc = docx.Document(path)
        parts: list[str] = []

        def _iter_block_items(parent):
            """Yield paragraphs and tables in document order."""
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            for child in parent.element.body:
                if child.tag == _qn("w:p"):
                    yield Paragraph(child, parent)
                elif child.tag == _qn("w:tbl"):
                    yield Table(child, parent)

        for block in _iter_block_items(doc):
            if hasattr(block, "text"):
                # Paragraph
                if block.text.strip():
                    parts.append(block.text)
            else:
                # Table — emit each row as tab-separated cells
                for row in block.rows:
                    cells = [c.text.strip() for c in row.cells]
                    row_text = "\t".join(cells)
                    if any(cells):
                        parts.append(row_text)

        # Also grab headers and footers
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf is not None:
                    for p in hf.paragraphs:
                        if p.text.strip():
                            parts.append(p.text)

        return "\n".join(parts)

    except ImportError:
        pass
    except Exception as e:
        return f"[DOCX read error: {e}]"

    # Fallback: unzip and strip XML tags from word/document.xml
    try:
        with open(path, "rb") as f:
            raw = f.read()
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if "word/document.xml" in z.namelist():
                xml = z.read("word/document.xml").decode("utf-8", errors="replace")
                text = re.sub(r"<[^>]+>", " ", xml)
                text = re.sub(r"\s+", " ", text).strip()
                return text
    except Exception as e:
        return f"[DOCX fallback read error: {e}]"

    return "[DOCX extraction requires python-docx: pip install python-docx]"


def _read_image(path: str, gemini_api_key: str = "") -> str:
    """
    Extract all text and describe the contents of an image document using
    Gemini's vision capability.  Returns a placeholder if no API key is set.
    """
    if not gemini_api_key:
        return "[Image file — set GOOGLE_API_KEY to enable vision extraction]"

    try:
        import google.generativeai as genai
        from google.generativeai.types import GenerationConfig

        genai.configure(api_key=gemini_api_key)
        uploaded = genai.upload_file(path=path)

        model = genai.GenerativeModel("gemini-flash-latest")
        response = model.generate_content(
            [
                uploaded,
                "This is a legal case document image. Do the following:\n"
                "1. Extract ALL visible text verbatim, preserving layout as much as possible.\n"
                "2. Describe any signatures, stamps, seals, handwriting, or annotations.\n"
                "3. Describe any tables, charts, or structured data.\n"
                "4. Note any dates, case numbers, names, or official markings.\n"
                "Be exhaustive — omit nothing.",
            ],
            generation_config=GenerationConfig(temperature=0.0, max_output_tokens=16384),
        )

        extracted = (response.text or "").strip()
        try:
            genai.delete_file(uploaded.name)
        except Exception:
            pass

        return extracted or "[Image extraction returned empty result]"

    except Exception as e:
        return f"[Image extraction failed: {e}]"


# ═══════════════════════════════════════════════════════════════════════════════
# DATE EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

_MONTHS = {
    "january":1,"february":2,"march":3,"april":4,"may":5,"june":6,
    "july":7,"august":8,"september":9,"october":10,"november":11,"december":12,
    "jan":1,"feb":2,"mar":3,"apr":4,"jun":6,"jul":7,"aug":8,
    "sep":9,"oct":10,"nov":11,"dec":12,
}

# Ordered from most- to least-specific so the first match wins.
_DATE_PATTERNS: List[Tuple[str, str]] = [
    # ISO: 2024-01-15 or 2024/01/15
    (r"\b(20\d{2})[-/](0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])\b", "%Y-%m-%d"),
    # US numeric: 01/15/2024 or 01-15-2024
    (r"\b(0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])[-/](20\d{2})\b", "%m-%d-%Y"),
    # Month-name long: January 15, 2024 / Jan 15 2024
    (r"\b(January|February|March|April|May|June|July|August|September|October|November|December"
     r"|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+"
     r"(\d{1,2}),?\s+(20\d{2})\b", "month_name"),
    # Ordinal: 15th day of January, 2024
    (r"\b(\d{1,2})(?:st|nd|rd|th)\s+day\s+of\s+"
     r"(January|February|March|April|May|June|July|August|September|October|November|December"
     r"|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec),?\s+(20\d{2})\b", "ordinal"),
    # Year-only: 2024 (weakest — used as fallback)
    (r"\b(20[12]\d)\b", "year_only"),
]


def _parse_date_match(pattern_type: str, groups: tuple) -> Optional[datetime]:
    try:
        if pattern_type == "%Y-%m-%d":
            return datetime(int(groups[0]), int(groups[1]), int(groups[2]))
        if pattern_type == "%m-%d-%Y":
            return datetime(int(groups[2]), int(groups[0]), int(groups[1]))
        if pattern_type == "month_name":
            month = _MONTHS.get(groups[0].lower().rstrip("."))
            if month:
                return datetime(int(groups[2]), month, int(groups[1]))
        if pattern_type == "ordinal":
            month = _MONTHS.get(groups[1].lower())
            if month:
                return datetime(int(groups[2]), month, int(groups[0]))
        if pattern_type == "year_only":
            return datetime(int(groups[0]), 1, 1)
    except (ValueError, IndexError):
        pass
    return None


def _extract_date(text: str) -> Optional[datetime]:
    """Return the first recognisable date found in *text*, or None."""
    for pattern, ptype in _DATE_PATTERNS:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            result = _parse_date_match(ptype, m.groups())
            if result:
                return result
    return None


def _document_date(doc: "CaseDocument") -> datetime:
    """
    Best-effort document date for chronological sorting.
    Priority: date in filename → date in first 3 000 chars of text → file mtime.
    Returns datetime.max for unreadable documents (sink them to the end).
    """
    if doc.read_error and not doc.text_content:
        return datetime.max

    # 1. Try filename first (most reliable for court-filed documents)
    stem = Path(doc.filename).stem
    d = _extract_date(stem.replace("_", " ").replace("-", " "))
    if d:
        return d

    # 2. Try the opening text of the document
    if doc.text_content:
        d = _extract_date(doc.text_content[:3000])
        if d:
            return d

    # 3. Fall back to filesystem mtime
    try:
        mtime = Path(doc.file_path).stat().st_mtime
        return datetime.fromtimestamp(mtime)
    except Exception:
        pass

    return datetime.max


# ═══════════════════════════════════════════════════════════════════════════════
# SCANNER
# ═══════════════════════════════════════════════════════════════════════════════

class CaseDirectoryScanner:
    """
    Scans a directory for legal documents and extracts their text content.

    Usage:
        scanner = CaseDirectoryScanner(gemini_api_key="...")
        documents = scanner.scan("/path/to/case/folder")
        stats = scanner.get_stats(documents)
    """

    def __init__(self, gemini_api_key: str = ""):
        self._gemini_api_key = gemini_api_key

    def scan(self, directory: str) -> List[CaseDocument]:
        """
        Scan the directory for all supported documents, extract text, and
        return them sorted chronologically by their document date.

        Raises FileNotFoundError / NotADirectoryError if the path is invalid.
        """
        dir_path = Path(directory)
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory}")

        file_paths = sorted(
            [f for f in dir_path.iterdir()
             if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS],
            key=lambda f: f.name.lower(),
        )
        return self._read_and_sort([str(f) for f in file_paths])

    def scan_files(self, file_paths: List[str]) -> List[CaseDocument]:
        """
        Read a list of individual file paths (not necessarily in one directory),
        extract text, and return them sorted chronologically by document date.
        Unsupported file types are silently skipped.
        """
        supported = [
            p for p in file_paths
            if Path(p).suffix.lower() in SUPPORTED_EXTENSIONS
        ]
        return self._read_and_sort(supported)

    def _read_and_sort(self, file_paths: List[str]) -> List[CaseDocument]:
        """Read files, respect MAX_TOTAL_CHARS, then sort chronologically."""
        documents: List[CaseDocument] = []
        total_chars = 0

        for path_str in file_paths:
            file_path = Path(path_str)
            doc = self._read_file(file_path)
            documents.append(doc)
            total_chars += doc.char_count

        # Sort chronologically; documents whose date cannot be determined go last
        documents.sort(key=_document_date)
        return documents

    def _read_file(self, file_path: Path) -> CaseDocument:
        ext = file_path.suffix.lower()
        doc_type = _infer_doc_type(file_path.name)

        try:
            if ext == ".pdf":
                text = _read_pdf(str(file_path), gemini_api_key=self._gemini_api_key)
            elif ext in (".docx", ".doc"):
                text = _read_docx(str(file_path))
            elif ext in IMAGE_EXTENSIONS:
                text = _read_image(str(file_path), gemini_api_key=self._gemini_api_key)
            else:
                text = _read_text_file(str(file_path))

            # Treat bracketed error strings from readers as failures, not content
            if text.startswith("[") and ("error" in text.lower() or "failed" in text.lower()):
                return CaseDocument(
                    filename=file_path.name,
                    file_path=str(file_path),
                    doc_type=doc_type,
                    text_content="",
                    char_count=0,
                    read_error=text.strip("[]"),
                )

            truncated = False  # No truncation

            return CaseDocument(
                filename=file_path.name,
                file_path=str(file_path),
                doc_type=doc_type,
                text_content=text,
                char_count=len(text),
                truncated=truncated,
            )

        except Exception as e:
            return CaseDocument(
                filename=file_path.name,
                file_path=str(file_path),
                doc_type=doc_type,
                text_content="",
                char_count=0,
                read_error=str(e),
            )

    def get_stats(self, documents: List[CaseDocument]) -> Dict:
        """Return summary statistics about the scanned documents."""
        readable  = [d for d in documents if not d.read_error and d.char_count > 0]
        errored   = [d for d in documents if d.read_error]
        by_type: Dict[str, int] = {}
        for d in readable:
            by_type[d.doc_type] = by_type.get(d.doc_type, 0) + 1
        return {
            "total":        len(documents),
            "readable":     len(readable),
            "errors":       len(errored),
            "by_type":      by_type,
            "total_chars":  sum(d.char_count for d in readable),
            "error_files":  [d.filename for d in errored],
        }

    def build_context_block(self, documents: List[CaseDocument]) -> str:
        """
        Format all documents into a single context block for AI prompts.

        Opens with a CHRONOLOGICAL CASE FILE INDEX (one line per document
        with its inferred date and type) so the agent immediately sees the
        timeline, then appends the full text of each document in order.
        """
        # ── Chronological index ───────────────────────────────────────────
        index_lines = ["CHRONOLOGICAL CASE FILE INDEX", "=" * 60]
        for i, doc in enumerate(documents, 1):
            date_obj = _document_date(doc)
            if date_obj == datetime.max:
                date_str = "date unknown"
            elif date_obj.month == 1 and date_obj.day == 1 and \
                    doc.text_content and str(date_obj.year) in doc.text_content[:3000]:
                # year-only match — don't show MM-DD
                date_str = str(date_obj.year)
            else:
                date_str = date_obj.strftime("%Y-%m-%d")

            status = f"[UNREADABLE: {doc.read_error}]" if doc.read_error else \
                     "[TRUNCATED]" if doc.truncated else ""
            index_lines.append(
                f"  {i:>2}. {date_str}  |  {doc.doc_type:<18}  |  "
                f"{doc.filename}  {status}".rstrip()
            )
        index_lines.append("=" * 60)
        index_block = "\n".join(index_lines)

        # ── Full document text ────────────────────────────────────────────
        parts = [index_block]
        for i, doc in enumerate(documents, 1):
            if doc.read_error or not doc.text_content:
                parts.append(
                    f"\n[DOCUMENT {i}: {doc.filename} | Type: {doc.doc_type}]\n"
                    f"[Could not read: {doc.read_error or 'empty file'}]\n"
                )
            else:
                truncation_note = " [TRUNCATED]" if doc.truncated else ""
                date_obj = _document_date(doc)
                date_label = date_obj.strftime("%Y-%m-%d") \
                    if date_obj != datetime.max else "date unknown"
                parts.append(
                    f"\n{'=' * 60}\n"
                    f"[DOCUMENT {i}: {doc.filename}"
                    f" | Type: {doc.doc_type}"
                    f" | Date: {date_label}"
                    f"{truncation_note}]\n"
                    f"{'=' * 60}\n"
                    f"{doc.text_content}\n"
                )
        return "\n".join(parts)
