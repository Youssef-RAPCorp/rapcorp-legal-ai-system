"""
Convert a plain-text legal document to a properly formatted .docx file
suitable for printing and filing in court.

Formatting:
  - Times New Roman 12pt (standard for court filings)
  - 1-inch margins on all sides
  - Double-spaced body paragraphs
  - Bold/centered headers
  - Page numbers (bottom center)
  - Certificate of Service section preserved

Usage:
    python scripts/txt_to_legal_docx.py --input output/.../motion.txt
    python scripts/txt_to_legal_docx.py --input output/.../verbal_guide.txt --style guide
"""

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    raise


def add_page_numbers(doc: Document) -> None:
    """Add centered page numbers to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        # Insert PAGE field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def set_paragraph_spacing(para, space_before=0, space_after=6, line_spacing=2.0):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE if line_spacing == 2.0 else WD_LINE_SPACING.SINGLE
    if line_spacing != 2.0:
        pf.line_spacing = Pt(line_spacing * 12)


def is_section_divider(line: str) -> bool:
    return bool(re.match(r'^[─━═\-=]{5,}', line.strip()))


def is_part_header(line: str) -> bool:
    """Lines like PART 1:, PART 2:, ── PART ..."""
    stripped = line.strip()
    return bool(re.match(r'^(PART\s+\d+[:\.]|──+\s*PART)', stripped, re.IGNORECASE))


def convert_motion_to_docx(text: str, output_path: Path) -> None:
    """Convert a court motion/filing text to a formatted DOCX."""
    doc = Document()

    # ── Page setup: 1-inch margins ──────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Default style: Times New Roman 12pt ──────────────────────────────
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    add_page_numbers(doc)

    lines = text.splitlines()
    i = 0
    in_caption = False  # Track the case caption / header block

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip pure divider lines
        if is_section_divider(stripped) and len(stripped) > 10:
            i += 1
            continue

        # Empty line → small gap
        if not stripped:
            p = doc.add_paragraph("")
            set_paragraph_spacing(p, space_before=0, space_after=0, line_spacing=1.0)
            i += 1
            continue

        # ── Case caption lines (IN THE ... COURT) ──────────────────────
        if stripped.startswith("IN THE") or stripped.startswith("IN RE"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, space_after=2, line_spacing=1.0)
            i += 1
            continue

        # ── Section headers (ALL CAPS lines that are standalone) ─────────
        if (stripped.isupper() and len(stripped) > 4
                and not stripped.startswith("WHEREFORE")
                and not stripped.startswith("COMES NOW")
                and not stripped.startswith("I HEREBY")):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.underline = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, space_before=12, space_after=6, line_spacing=1.0)
            i += 1
            continue

        # ── Numbered paragraphs (1. 2. 3. a. b.) ────────────────────────
        num_match = re.match(r'^(\d+\.|[a-z]\.|[a-z]\))\s+(.*)', stripped)
        if num_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            set_paragraph_spacing(p, space_after=6)
            i += 1
            continue

        # ── Signature / date lines ──────────────────────────────────────
        if stripped.startswith("_____") or stripped.startswith("DATED"):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.left_indent = Inches(0)
            run = p.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            set_paragraph_spacing(p, space_before=6, space_after=4, line_spacing=1.0)
            i += 1
            continue

        # ── Default body paragraph ───────────────────────────────────────
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        set_paragraph_spacing(p, space_after=6)
        i += 1

    doc.save(str(output_path))
    print(f"  Saved: {output_path}")


def convert_guide_to_docx(text: str, output_path: Path) -> None:
    """Convert the verbal argument guide to a readable, printable DOCX."""
    doc = Document()

    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_page_numbers(doc)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("VERBAL ARGUMENT GUIDE — SARPY COUNTY COURT\nIn re Guardianship of Youssef H. Eweis, Case No. PR 26-125")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_para, space_before=0, space_after=18, line_spacing=1.0)

    lines = text.splitlines()

    for line in lines:
        stripped = line.strip()

        if not stripped:
            p = doc.add_paragraph("")
            set_paragraph_spacing(p, space_before=0, space_after=2, line_spacing=1.0)
            continue

        # Section dividers / headers like ─────PART 1─────
        if is_section_divider(stripped) or is_part_header(stripped):
            # The actual part title is the next non-divider line — handle inline
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("─━═ \t"))
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)  # Navy blue for headers
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, space_before=14, space_after=4, line_spacing=1.0)
            continue

        # Bold **text** markers
        if "**" in stripped:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*', stripped)
            bold_on = False
            for part in parts:
                if part:
                    run = p.add_run(part)
                    run.bold = bold_on
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                bold_on = not bold_on
            set_paragraph_spacing(p, space_before=4, space_after=4, line_spacing=1.0)
            continue

        # Bullet points
        if stripped.startswith("* ") or stripped.startswith("- ") or stripped.startswith("• "):
            content = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            run = p.add_run(f"• {content}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            set_paragraph_spacing(p, space_before=2, space_after=2, line_spacing=1.0)
            continue

        # *Italicised stage directions* like *[CRITICAL TIMING: ...]*
        if stripped.startswith("*[") or stripped.startswith("*(*") or (stripped.startswith("*") and stripped.endswith("*")):
            p = doc.add_paragraph()
            clean = stripped.strip("*").strip()
            run = p.add_run(clean)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            set_paragraph_spacing(p, space_before=6, space_after=4, line_spacing=1.0)
            continue

        # Quoted speech (lines in quotes)
        if stripped.startswith('"') and stripped.endswith('"'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(stripped)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            set_paragraph_spacing(p, space_before=6, space_after=6, line_spacing=1.0)
            continue

        # Default body
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        set_paragraph_spacing(p, space_before=2, space_after=4, line_spacing=1.0)

    doc.save(str(output_path))
    print(f"  Saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Convert legal .txt to formatted .docx")
    parser.add_argument("--input", required=True, help="Path to the .txt file to convert")
    parser.add_argument("--style", default="motion",
                        choices=["motion", "guide"],
                        help="'motion' for court filings, 'guide' for verbal argument guides")
    args = parser.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        print(f"ERROR: File not found: {in_path}")
        return

    text = in_path.read_text(encoding="utf-8")
    out_path = in_path.with_suffix(".docx")

    print(f"\n  Converting: {in_path.name} → {out_path.name}")

    if args.style == "guide":
        convert_guide_to_docx(text, out_path)
    else:
        convert_motion_to_docx(text, out_path)

    print(f"  Done. Open and print: {out_path}")


if __name__ == "__main__":
    main()
